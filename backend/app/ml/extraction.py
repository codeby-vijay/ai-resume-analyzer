"""Resume text extraction from PDF and DOCX files."""
import io
import pdfplumber
from docx import Document as DocxDocument
from fastapi import HTTPException


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Could not parse PDF: {exc}")
    return "\n".join(text_parts).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs).strip()
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Could not parse DOCX: {exc}")


def extract_resume_text(filename: str, file_bytes: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif lower.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    else:
        raise HTTPException(
            status_code=400, detail="Unsupported file type. Please upload a PDF or DOCX file."
        )
    if not text or len(text.strip()) < 30:
        raise HTTPException(
            status_code=400,
            detail="Could not extract readable text from this resume. Try a different file.",
        )
    return text
